# report_generator.py
# SpectralG — Clinical Report Generator v2.1
# Produces research-grade reports with full ACMG table, HGVS, evidence panel
#
# Framework: ACMG/AMP 2015 + 2023 updates | PP5 not applied
# Anti-hallucination: temperature=0.1, constrained prompt, disclaimer block

import os
import re
import json
import anthropic
from dotenv import load_dotenv
from pathlib import Path
from datetime import date

load_dotenv(dotenv_path=Path.home() / ".env", override=True)
load_dotenv()


# ── CLINICAL INTERPRETATION LOGIC ────────────────────────────────────────────

def get_clinical_interpretation(variant):
    """
    Rule-based clinical interpretation for a single variant.
    Used in both AI prompt and fallback report.
    Returns plain English interpretation string.
    """
    acmg        = variant.get("acmg", "VUS")
    priority    = variant.get("priority", "LOW")
    gene        = variant.get("gene", "Unknown")
    confidence  = variant.get("confidence_level", "Limited")

    if acmg == "Pathogenic":
        return (f"This variant in {gene} is classified as Pathogenic "
                f"(confidence: {confidence}). It is likely disease-causing. "
                "Clinical correlation and cascade family testing are recommended.")

    elif acmg == "Likely Pathogenic":
        return (f"This variant in {gene} is classified as Likely Pathogenic "
                f"(confidence: {confidence}). It is probably disease-associated. "
                "Functional validation and parental testing are recommended.")

    elif acmg == "VUS":
        if priority == "HIGH":
            return (f"This variant in {gene} is a Variant of Uncertain Significance "
                    f"(VUS) with HIGH computational priority score "
                    f"(confidence: {confidence}). "
                    "Clinical significance cannot be established without additional evidence. "
                    "Parental testing and annual reclassification review are recommended. "
                    "Do not make clinical decisions based on VUS classification alone.")
        else:
            return (f"This variant in {gene} is a Variant of Uncertain Significance "
                    f"(VUS, confidence: {confidence}). "
                    "Clinical significance cannot be established from current evidence. "
                    "Annual reclassification review is recommended.")

    elif acmg == "Likely Benign":
        return (f"This variant in {gene} is classified as Likely Benign "
                f"(confidence: {confidence}). "
                "It is unlikely to be the primary cause of disease in isolation.")

    elif acmg == "Benign":
        return (f"This variant in {gene} is classified as Benign "
                f"(confidence: {confidence}). "
                "It is a common population variant and unlikely to cause disease.")

    return (f"Classification for {gene} variant could not be established. "
            "Manual review required.")


def format_gnomad(gnomad_af):
    """Format gnomAD frequency for report text."""
    if gnomad_af is None:
        return "Not available in gnomAD (variant may be novel or rare)"
    if isinstance(gnomad_af, dict):
        sas   = gnomad_af.get("south_asian")
        glb   = gnomad_af.get("global")
        parts = []
        if sas is not None:
            parts.append(f"South Asian (SAS): {sas:.6f}")
        if glb is not None:
            parts.append(f"Global: {glb:.6f}")
        return " | ".join(parts) if parts else "Not available in VEP response"
    try:
        return f"{float(gnomad_af):.6f}"
    except Exception:
        return "Not available"


def format_evidence_list(criteria_table):
    """
    Summarise applied ACMG criteria for report text.
    Returns human-readable string.
    """
    if not criteria_table:
        return "ACMG criteria table not available."

    applied     = [c for c in criteria_table
                   if c.get("applied") and c["code"] not in {"PP5", "BP6"}]
    not_applied = [c for c in criteria_table
                   if not c.get("applied") and c["code"] not in {"PP5", "BP6"}]

    lines = []
    if applied:
        lines.append(f"Applied criteria ({len(applied)}): " +
                     ", ".join(c["code"] for c in applied))
    else:
        lines.append("No pathogenic or benign criteria applied.")

    lines.append(f"Not applied: {len(not_applied)} criteria "
                 f"(see full criteria table in Section 5).")
    lines.append("PP5 intentionally excluded per ACMG 2023 guidance.")
    return " | ".join(lines)


# ── BUILD VARIANT BLOCK ───────────────────────────────────────────────────────

def build_variant_block(v, index):
    """
    Build a complete structured dict for one variant.
    Used in both AI prompt and fallback report.
    """
    ann = v.get("annotation", {})
    return {
        "index":          index + 1,
        "gene":           v.get("gene", "Unknown"),
        "position":       f"chr{v.get('chrom','?')}:{v.get('pos','?')}",
        "change":         f"{v.get('ref','?')} > {v.get('alt','?')}",
        "hgvs_c":         v.get("hgvsc", "Not available"),
        "hgvs_p":         v.get("hgvsp", "Not available"),
        "consequence":    ann.get("consequence", "unknown"),
        "impact":         ann.get("impact", "UNKNOWN"),
        "sift":           str(ann.get("sift", "N/A")),
        "polyphen":       str(ann.get("polyphen", "N/A")),
        "gnomad_af":      format_gnomad(v.get("gnomad_af")),
        "clinvar":        v.get("clinvar", "Unknown"),
        "acmg_class":     v.get("acmg", "VUS"),
        "confidence":     v.get("confidence_level", "Limited"),
        "priority":       v.get("priority", "LOW"),
        "acmg_evidence":  format_evidence_list(v.get("acmg_criteria_table", [])),
        "interpretation": get_clinical_interpretation(v),
    }


# ── HALLUCINATION CHECKER ─────────────────────────────────────────────────────

def check_for_hallucinations(report_text, variant_summary):
    """
    Scan AI report for common hallucination patterns.
    Returns (is_safe, warnings) tuple.
    """
    warnings     = []
    report_lower = report_text.lower()

    # Invented statistics
    stat_patterns = [
        r'\d+%\s+of\s+(patients|cases|individuals)',
        r'reported\s+in\s+\d+\s+(families|patients|cases)',
        r'prevalence\s+of\s+\d+',
        r'affects\s+\d+\s+in\s+\d+',
    ]
    for pattern in stat_patterns:
        if re.search(pattern, report_lower):
            warnings.append(
                f"⚠️ Possible invented statistic detected. "
                f"Verify against ClinVar before clinical reporting."
            )

    # Invented citations
    citation_patterns = [
        r'\(pmid\s*:\s*\d+\)',
        r'et al\.\s*\(\d{4}\)',
    ]
    for pattern in citation_patterns:
        if re.search(pattern, report_lower):
            warnings.append(
                "⚠️ Possible invented citation detected. "
                "Verify any references manually."
            )

    # ClinVar claims without source data
    clinvar_in_data = any(
        v.get("clinvar", "Unknown").lower() not in ("unknown", "")
        for v in variant_summary
    )
    if not clinvar_in_data and "clinvar" in report_lower:
        if "pathogenic" in report_lower or "benign" in report_lower:
            warnings.append(
                "⚠️ Report mentions ClinVar classification but "
                "source data has no ClinVar data. Verify manually."
            )

    # Less common genes — flag for OMIM verification
    known_safe_genes = {
        "TP53","BRCA1","BRCA2","CFTR","HBB","MLH1",
        "MSH2","APC","PTEN","PKD1","PKD2",
        "MYBPC3","MYH7","KCNQ1","SCN5A","SCN1A",
        "GJB2","DPYD","CYP2D6","CYP2C19","TPMT"
    }
    report_genes = {v.get("gene","Unknown") for v in variant_summary
                   if v.get("gene","Unknown") not in ("Unknown","")}
    unknown_genes = report_genes - known_safe_genes
    if unknown_genes:
        warnings.append(
            f"ℹ️ Report includes less common genes: {unknown_genes}. "
            f"Verify gene function descriptions against OMIM."
        )

    return len(warnings) == 0, warnings


# ── AI REPORT GENERATION ──────────────────────────────────────────────────────

def generate_report(variants, patient_id="SAMPLE_001",
                    clinical_info=None, language="English"):
    """
    Generate AI clinical report using Claude.
    Falls back to rule-based report if API unavailable or fails.

    Args:
        variants:      list of annotated variant dicts
        patient_id:    sample/report ID string
        clinical_info: dict from input_handler.parse_clinical_info()
        language:      report language string
    """
    if clinical_info is None:
        clinical_info = {}

    today = date.today().strftime("%d %B %Y")

    # ── Build variant summaries ───────────────────────────────────────────────
    variant_blocks = [build_variant_block(v, i) for i, v in enumerate(variants)]

    # ── Extract clinical context ──────────────────────────────────────────────
    sex         = clinical_info.get("sex", "Sex: not provided in referral")
    indication  = clinical_info.get("indication", "Not provided")
    report_type = clinical_info.get("report_type", "Clinical WES")
    gp          = clinical_info.get("genotype_phenotype_correlation", "Not assessed")
    features    = clinical_info.get("clinical_features", "Not provided")
    referring   = clinical_info.get("referring_clinician", "Not provided")

    # ── High-priority variants for executive summary ──────────────────────────
    high_priority = [v for v in variants if v.get("priority") == "HIGH"]
    pathogenic    = [v for v in variants if v.get("acmg") in
                     ("Pathogenic", "Likely Pathogenic")]

    # ── Build single clean prompt ─────────────────────────────────────────────
    prompt = f"""You are a clinical genetics report writer for SpectralG.

CRITICAL ANTI-HALLUCINATION RULES — FOLLOW STRICTLY:

1. ONLY use data provided in the variant JSON below.
   Never invent gene functions, disease associations, or frequencies.

2. If a field says "Not available", "Unknown", or "N/A" —
   write "Not available" in the report. Never fill gaps with guesses.

3. Never cite specific papers, PMIDs, or studies unless explicitly
   provided in the data below.

4. Never state specific patient numbers, prevalence statistics,
   or population frequencies unless they are in the data below.

5. Never say a variant "has been reported in X families" or
   "causes disease in X% of cases" unless the data states this.

6. For gene function — use only well-established textbook-level facts.
   If unsure — write "Gene function: refer to OMIM [gene]."

7. For ACMG classification — only state criteria marked applied=True
   in the criteria table. Never add criteria not in the table.

8. Never say "diagnosis confirmed" — use "findings suggest" or
   "consistent with".

9. Never infer sex from name — use only what was provided.

10. End every variant interpretation with:
    "This interpretation is based solely on computational evidence
    provided by SpectralG. Manual review against ClinVar, Franklin,
    and published literature is required before clinical reporting."

Report type: {report_type}
Patient ID: {patient_id}
Sex: {sex}
Age: {clinical_info.get('age','Not provided')}
Referring clinician: {referring}
Clinical indication: {indication}
Clinical features: {features}
Genotype-phenotype correlation: {gp}
Total variants analysed: {len(variants)}
High-priority variants: {len(high_priority)}
Pathogenic/Likely Pathogenic: {len(pathogenic)}

VARIANT DATA (use ONLY this):
{json.dumps(variant_blocks, indent=2)}

Generate a structured clinical report with these exact sections:

## 1. EXECUTIVE SUMMARY
- 2-3 sentences summarising the key finding
- State total variants and how many are HIGH priority
- State genotype-phenotype correlation explicitly
- Highlight any Pathogenic or Likely Pathogenic findings

## 2. VARIANT INTERPRETATIONS
For EACH variant:
- Gene name and full genomic position
- HGVS c. and HGVS p. notation (state "Not available" if missing)
- ACMG classification with confidence level
- Applied ACMG evidence criteria
- Plain language clinical relevance (2-3 sentences)
- Priority level with justification

## 3. CLINICAL RECOMMENDATIONS
- Specific recommendations based on findings
- Use "Standard-of-care per guideline, to be implemented by treating team"
- Never say treatment is mandatory or not optional
- Suggest: confirmatory testing, specialist referral, cascade family testing
- State "Do not make clinical decisions based on VUS alone" where applicable

## 4. GENOTYPE-PHENOTYPE CORRELATION
- State: Present / Absent / Partial / Not assessed
- Brief explanation

## 5. METHODS NOTE
- SpectralG used Ensembl VEP annotation, ACMG/AMP 2015 + 2023 criteria,
  gnomAD South Asian subpopulation, PP5 not applied per ACMG 2023

## 6. DISCLAIMER
- AI-assisted research tool — requires qualified clinical geneticist review
- Not for clinical use without professional validation
- PP5 not applied per ACMG 2023

Rules:
- Keep total report under 700 words
- Generate entirely in {language}
- Keep gene names, HGVS notation, ACMG codes in English regardless of language
"""

    # ── Try Claude AI at temperature 0.1 (factual mode) ──────────────────────
    ai_text = None
    try:
        api_key = os.getenv("ANTHROPIC_API_KEY")
        if api_key:
            client   = anthropic.Anthropic(api_key=api_key)
            response = client.messages.create(
                model       = "claude-haiku-4-5-20251001",
                max_tokens  = 1200,
                temperature = 0.1,   # Factual mode — prevents hallucination
                messages    = [{"role": "user", "content": prompt}]
            )
            ai_text = response.content[0].text
            print("✅ AI report generated (temperature=0.1)")
        else:
            print("⚠️ No API key — using rule-based fallback")
    except Exception as e:
        print(f"AI error: {e}")

    # ── Hallucination check ───────────────────────────────────────────────────
    if ai_text:
        is_safe, warnings = check_for_hallucinations(ai_text, variant_blocks)
        if warnings:
            warning_block = "\n\n---\n⚠️ AUTOMATED QUALITY FLAGS:\n"
            warning_block += "\n".join(warnings)
            warning_block += ("\n\nVerify flagged items manually before "
                              "clinical reporting.\n---\n")
            ai_text += warning_block

    # ── Mandatory disclaimer appended to every AI report ─────────────────────
    disclaimer_block = f"""

---
⚠️ AI REPORT QUALITY NOTICE

This clinical interpretation was generated by Claude AI (Anthropic)
using only the variant data provided by SpectralG.

Before clinical use — manually verify:
□ Gene function against OMIM (omim.org)
□ Variant classification against ClinVar (ncbi.nlm.nih.gov/clinvar)
□ Population frequency against gnomAD (gnomad.broadinstitute.org)
□ Any statistics or patient numbers cited
□ Any disease associations mentioned

AI models can generate plausible-sounding but incorrect statements.
SpectralG is a research tool — all reports require review by a
qualified clinical geneticist before clinical use.

PP5 not applied per ACMG 2023 (Biesecker & Harrison).
Report generated: {today}
---
"""

    if ai_text:
        ai_text += disclaimer_block
        return ai_text

    # ── Fallback to rule-based report ─────────────────────────────────────────
    return generate_rule_based_report(variants, patient_id, clinical_info)


# ── RULE-BASED FALLBACK REPORT ────────────────────────────────────────────────

def generate_rule_based_report(variants, patient_id="SAMPLE_001",
                                clinical_info=None):
    """
    Generate a complete rule-based report without AI.
    Includes full variant details, ACMG criteria, and evidence.
    Used when Claude API is unavailable.
    """
    if clinical_info is None:
        clinical_info = {}

    today  = date.today().strftime("%d %B %Y")
    high_p = [v for v in variants if v.get("priority") == "HIGH"]
    path_v = [v for v in variants if v.get("acmg") in
               ("Pathogenic", "Likely Pathogenic")]

    report = f"""# SpectralG Clinical Variant Report
## {patient_id} | {today}

---

## 1. EXECUTIVE SUMMARY

{len(variants)} variant(s) analysed from {clinical_info.get('report_type','clinical sequencing')}.
{len(high_p)} high-priority variant(s) detected.
{len(path_v)} Pathogenic or Likely Pathogenic variant(s) identified.
Genotype-phenotype correlation: {clinical_info.get('genotype_phenotype_correlation','Not assessed')}.

{'⚠️ Pathogenic or Likely Pathogenic variants identified — clinical review required.' if path_v else 'No Pathogenic or Likely Pathogenic variants identified in this dataset.'}

---

## 2. VARIANT INTERPRETATIONS

"""
    for i, v in enumerate(variants):
        block            = build_variant_block(v, i)
        ct               = v.get("acmg_criteria_table", [])
        applied_criteria = [c["code"] for c in ct
                            if c.get("applied") and
                            c["code"] not in {"PP5", "BP6"}]

        report += f"""### Variant {i+1}: {block['gene']}

| Field | Value |
|-------|-------|
| Gene | {block['gene']} |
| Position | {block['position']} |
| Change | {block['change']} |
| HGVS c. | {block['hgvs_c']} |
| HGVS p. | {block['hgvs_p']} |
| Consequence | {block['consequence']} |
| VEP Impact | {block['impact']} |
| SIFT | {block['sift']} |
| PolyPhen-2 | {block['polyphen']} |
| gnomAD AF | {block['gnomad_af']} |
| ClinVar | {block['clinvar']} |
| ACMG Classification | {block['acmg_class']} |
| Evidence Confidence | {block['confidence']} |
| Priority | {block['priority']} |
| Applied Criteria | {', '.join(applied_criteria) if applied_criteria else 'None'} |

**Interpretation:** {block['interpretation']}

**Note on PP5:** PP5 not applied per ACMG 2023 guidance (Biesecker & Harrison).
ClinVar data documented for reference but not counted as independent evidence.

---

"""

    report += "## 3. CLINICAL RECOMMENDATIONS\n\n"

    if path_v:
        report += (
            "- Confirm all Pathogenic/Likely Pathogenic variants by Sanger "
            "sequencing before cascade family testing.\n"
            "- Refer to relevant specialist — standard-of-care per guideline, "
            "to be implemented by the treating team.\n"
            "- Offer genetic counselling to discuss inheritance, recurrence risk, "
            "and reproductive options.\n"
        )
    elif high_p:
        report += (
            "- HIGH priority VUS variants warrant further investigation.\n"
            "- Parental testing recommended to assess de novo vs inherited status "
            "— may upgrade VUS classification.\n"
            "- Annual reclassification review recommended as databases update.\n"
            "- Do not make clinical management decisions based on VUS alone.\n"
        )
    else:
        report += (
            "- No Pathogenic or Likely Pathogenic variants identified.\n"
            "- VUS findings should be re-evaluated annually.\n"
            "- Clinical correlation with presenting phenotype is essential.\n"
        )

    report += f"""
## 4. GENOTYPE-PHENOTYPE CORRELATION

Status: {clinical_info.get('genotype_phenotype_correlation','Not assessed')}

{clinical_info.get('gp_narrative',
  'Genotype-phenotype correlation could not be assessed — clinical features '
  'not provided. The ordering clinician should assess whether identified '
  'variants are consistent with the clinical presentation.')}

---

## 5. METHODS NOTE

Variant annotation: Ensembl VEP REST API (GRCh38/hg38).
Gene identification fallback: Ensembl Overlap API.
Population frequencies: gnomAD v4.1 — South Asian (SAS) subpopulation prioritised for Indian patients.
Classification framework: ACMG/AMP 2015 (Richards et al., Genet Med 2015) + 2023 updates (Biesecker & Harrison, Genet Med 2023) + ACGS Best Practice Guidelines v4.1 (2024).
PP3 at supporting level only per Pejaver et al. 2022.
PP5 not applied per ACMG 2023 guidance.
REVEL threshold ≥0.733 recommended for PP3 (not available via VEP REST API — manual lookup required).
Computational tools: SIFT, PolyPhen-2.
Database access date: {today}.

---

## 6. DISCLAIMER

This report was generated by SpectralG (Prakash NK, MSc Human Genetics),
an AI-assisted research tool for scientific interpretation support.

This report does NOT constitute a clinical diagnostic test result, medical
diagnosis, or clinical advice. All findings require review and validation
by a qualified clinical geneticist before any clinical use.

Variant classifications may change as evidence accumulates.
VUS findings should be re-evaluated annually.
PP5 has not been applied per ACMG 2023.

Curation date: {today} | SpectralG | Prakash NK | Hyderabad, India
"""
    return report


# ── WORD DOCUMENT GENERATION ──────────────────────────────────────────────────

def generate_word_report(variants, report_text, patient_id="SAMPLE_001",
                          language="English"):
    """
    Generate a formatted Word (.docx) report.
    Includes: variant table, ACMG criteria table, full report text.
    Returns filename of saved .docx.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading(
        "SpectralG — Clinical Variant Interpretation Report", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f"Sample ID:      {patient_id}")
    doc.add_paragraph(f"Report date:    {date.today().strftime('%d %B %Y')}")
    doc.add_paragraph(f"Language:       {language}")
    doc.add_paragraph(f"Framework:      ACMG/AMP 2015 + 2023 | PP5 not applied")
    doc.add_paragraph(f"Tool:           SpectralG v2.1 | Ensembl VEP + Claude AI")
    doc.add_paragraph(f"Variants:       {len(variants)}")
    doc.add_paragraph("")

    # Variant summary table
    doc.add_heading("Variant Summary", level=1)
    tbl = doc.add_table(rows=1, cols=8)
    tbl.style = "Table Grid"
    for i, h in enumerate(["Gene","Position","HGVS c.","HGVS p.",
                            "Consequence","ACMG","Confidence","Priority"]):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for v in variants:
        ann = v.get("annotation", {})
        row = tbl.add_row().cells
        row[0].text = v.get("gene", "Unknown")
        row[1].text = f"chr{v.get('chrom','')}:{v.get('pos','')}"
        row[2].text = v.get("hgvsc", "N/A")
        row[3].text = v.get("hgvsp", "N/A")
        row[4].text = ann.get("consequence", "unknown")
        row[5].text = v.get("acmg", "VUS")
        row[6].text = v.get("confidence_level", "N/A")
        row[7].text = v.get("priority", "LOW")

    doc.add_paragraph("")

    # ACMG criteria table per variant
    for i, v in enumerate(variants):
        ct = v.get("acmg_criteria_table", [])
        if not ct:
            continue
        doc.add_heading(
            f"ACMG Criteria — {v.get('gene','Unknown')} "
            f"(chr{v.get('chrom','')}:{v.get('pos','')})",
            level=2
        )
        doc.add_paragraph(
            f"Classification: {v.get('acmg','VUS')} | "
            f"Confidence: {v.get('confidence_level','N/A')} | "
            f"PP5 not applied (ACMG 2023)"
        )
        ct_tbl = doc.add_table(rows=1, cols=3)
        ct_tbl.style = "Table Grid"
        for j, h in enumerate(["Criterion", "Weight", "Applied"]):
            ct_tbl.rows[0].cells[j].text = h
            for run in ct_tbl.rows[0].cells[j].paragraphs[0].runs:
                run.bold = True
        for c in ct:
            r = ct_tbl.add_row().cells
            r[0].text = c["code"]
            r[1].text = c.get("weight", "")
            r[2].text = "✓ YES" if c.get("applied") else "No"
        doc.add_paragraph("")

    # Full report text
    doc.add_heading("Clinical Interpretation", level=1)
    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)
        else:
            doc.add_paragraph(line)

    # Disclaimer
    doc.add_heading("Disclaimer", level=1)
    p = doc.add_paragraph(
        "SpectralG is an AI-assisted research tool. "
        "All reports require review by a qualified clinical geneticist "
        "before clinical use. PP5 not applied per ACMG 2023 guidance. "
        f"Report date: {date.today().strftime('%d %B %Y')} | "
        "Prakash NK | SpectralG | Hyderabad, India"
    )
    for run in p.runs:
        run.italic = True

    filename = f"{patient_id}_SpectralG_report.docx"
    doc.save(filename)
    print(f"✅ Word report saved: {filename}")
    return filename